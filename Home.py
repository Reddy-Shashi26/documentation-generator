import streamlit as st
import os
from langchain.globals import set_verbose
from langchain_groq import ChatGroq
from langchain_google_genai import ChatGoogleGenerativeAI
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import io

# Set theme and configuration
st.set_page_config(
    layout="wide",
    page_title="Advanced Documentation Generator",
    initial_sidebar_state="expanded"
)

# Add custom CSS for dark mode
st.markdown("""
<style>
:root {
    --primary-color: #ff4b4b;
}
.dark-mode {
    background-color: #1e1e1e;
    color: #ffffff;
}
</style>
""", unsafe_allow_html=True)

# Initialize session states
if 'dark_mode' not in st.session_state:
    st.session_state.dark_mode = False
if 'template_type' not in st.session_state:
    st.session_state.template_type = 'technical'
if 'autosave' not in st.session_state:
    st.session_state.autosave = True
if 'doc_sections' not in st.session_state:
    st.session_state.doc_sections = {}
if 'metadata' not in st.session_state:
    st.session_state.metadata = {}
if 'feedback_history' not in st.session_state:
    st.session_state.feedback_history = []
if 'has_generated' not in st.session_state:
    st.session_state.has_generated = False

# Sidebar configuration
with st.sidebar:
    st.title("Settings")
    st.session_state.dark_mode = st.toggle("Dark Mode", st.session_state.dark_mode)
    st.session_state.template_type = st.selectbox(
        "Document Template",
        ['technical', 'research', 'business', 'academic']
    )
    st.session_state.autosave = st.toggle("Auto-save", st.session_state.autosave)

# Set verbosity
set_verbose(False)

# Load environment variables
load_dotenv()

# Load API keys
groq_api_key = os.getenv('GROQ_API_KEY')
os.environ["GOOGLE_API_KEY"] = os.getenv("GOOGLE_API_KEY")

# Initialize AI models
groq_model = ChatGroq(groq_api_key=groq_api_key, model_name="openai/gpt-oss-120b")
google_model = ChatGoogleGenerativeAI(model="gemini-2.5-flash")

# Define documentation sections with updated structure
SECTIONS = {
    "COVER_PAGE": {"title": "COVER PAGE", "pages": 1},
    "CONTENTS": {"title": "TABLE OF CONTENTS", "pages": 1},
    "ABSTRACT": {"title": "ABSTRACT", "pages": 1, "description": "Provides a high-level summary of the project, including objectives, methodology, and expected outcomes."},
    "LITERATURE_SURVEY": {"title": "LITERATURE SURVEY", "pages": 3, "description": "Reviews existing research, technologies, and methodologies related to the project."},
    "CHAPTER_1": {"title": "CHAPTER 1: INTRODUCTION", "pages": 3, "description": "Explains the background, significance, and objectives of the project."},
    "CHAPTER_2": {"title": "CHAPTER 2: WORKING AND BLOCK DIAGRAM", "pages": 4, "description": "Details the working principle and block diagram representation of the system."},
    "CHAPTER_3": {"title": "CHAPTER 3: HARDWARE COMPONENTS DESCRIPTION", "pages": 5, "description": "Describes the components used (Arduino, Bluetooth module, motor driver, sensors, etc.)."},
    "CHAPTER_4": {"title": "CHAPTER 4: SOFTWARE DESCRIPTION", "pages": 4, "description": "Covers programming aspects, including Arduino code and mobile app interface."},
    "CHAPTER_5": {"title": "CHAPTER 5: ADVANTAGES & APPLICATIONS", "pages": 2, "description": "Highlights benefits and potential real-world applications of the project."},
    "CHAPTER_6": {"title": "CHAPTER 6: FUTURE SCOPE & CONCLUSION", "pages": 2, "description": "Discusses improvements, extensions, and final remarks on the project."},
    "REFERENCES": {"title": "REFERENCES", "pages": 1, "description": "Lists cited sources, books, research papers, and online materials."}
}

def setup_document_styles(doc):
    """Set up professional document styles with updated formatting"""
    # Default paragraph style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)  # Black color
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph_format.space_after = Pt(12)

    # Heading styles with borders
    heading_sizes = {
        'Heading 1': 16,
        'Heading 2': 14,
        'Heading 3': 13,
        'Heading 4': 12
    }
    
    for style_name, size in heading_sizes.items():
        style = doc.styles[style_name]
        font = style.font
        font.name = 'Arial'
        font.size = Pt(size)
        font.bold = True
        font.color.rgb = RGBColor(0, 0, 0)  # Black color
        font.all_caps = True  # Make headings all caps
        paragraph_format = style.paragraph_format
        paragraph_format.space_before = Pt(24)
        paragraph_format.space_after = Pt(12)
        paragraph_format.keep_with_next = True
        # The next two lines add borders to headings
        paragraph_format.border_bottom = True
        paragraph_format.border_top = True

def add_placeholder_toc(doc):
    """Add a placeholder for table of contents"""
    doc.add_heading("TABLE OF CONTENTS", 1)
    # Add placeholder text for TOC
    p = doc.add_paragraph()
    p.add_run("This Table of Contents will automatically update when you open the document in Microsoft Word. Right-click and select 'Update Field' to update it.")
    
    # Add instructions
    p = doc.add_paragraph()
    p.add_run("To replace this with an actual Table of Contents in Word:").bold = True
    p = doc.add_paragraph("1. Delete this text")
    p = doc.add_paragraph("2. Go to References tab")
    p = doc.add_paragraph("3. Click 'Table of Contents'")
    p = doc.add_paragraph("4. Select a style")

def create_docx(metadata):
    """Create a professionally formatted DOCX document with updated structure"""
    doc = Document()
    setup_document_styles(doc)
    
    # Set margins
    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        # Add page numbers by setting different first page
        section.different_first_page_header_footer = True
    
    # Cover Page
    cover = doc.add_heading(metadata.get('title', 'PROJECT DOCUMENTATION').upper(), 0)
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata
    for field in ['author', 'institution', 'date']:
        if metadata.get(field):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(metadata[field].upper())
            run.font.size = Pt(12)
    
    # Add space for image
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("[Space reserved for project image]")
    
    doc.add_page_break()
    
    # Add table of contents
    add_placeholder_toc(doc)
    doc.add_page_break()
    
    # Add content sections with proper page breaks and borders
    for section_key, section_info in SECTIONS.items():
        # Skip cover page and contents as they're already added
        if section_key in ['COVER_PAGE', 'CONTENTS']:
            continue
            
        # Add section heading
        heading = doc.add_heading(section_info['title'], 1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add section content
        if section_key in st.session_state.doc_sections:
            content = st.session_state.doc_sections[section_key]
            paragraphs = content.split('\n\n')
            
            for para in paragraphs:
                if para.strip():
                    if para.strip().startswith('#'):
                        # Handle subsection headings
                        level = len(para.split()[0].strip('#'))
                        text = ' '.join(para.split()[1:]).upper()
                        subheading = doc.add_heading(text, level + 1)
                        subheading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    else:
                        # Regular paragraphs
                        p = doc.add_paragraph(para.strip())
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        # Add page break after each section
        doc.add_page_break()
    
    # Add footer with note about page numbers
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Page ")
    p.add_run("- {PAGE} -")
    p.add_run(" (Open in Word to see actual page numbers)")
    
    # Save document
    doc_path = "project_documentation.docx"
    doc.save(doc_path)
    return doc_path

# Main UI layout
def main():
    nav_col1, nav_col2, nav_col3 = st.columns([1, 2, 1])

    with nav_col2:
        st.title("Advanced Project Documentation Generator")

    with nav_col3:
        if st.session_state.has_generated and os.path.exists("project_documentation.docx"):
            with open("project_documentation.docx", "rb") as file:
                st.download_button(
                    label="📥 Download Documentation",
                    data=file,
                    file_name="project_documentation.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

    # Main layout
    col1, col2 = st.columns([2, 3])

    with col1:
        st.subheader("Project Information")
        
        # Project metadata
        with st.expander("Project Details", expanded=True):
            project_title = st.text_input("Project Title", key="title")
            author_name = st.text_input("Author Name", key="author")
            institution = st.text_input("Institution Name", key="institution")
            project_type = st.selectbox(
                "Project Type",
                ["Arduino-based", "Raspberry Pi-based", "IoT Project", "Embedded Systems", "Other"],
                key="project_type"
            )
            st.session_state.metadata = {
                "title": project_title,
                "author": author_name,
                "institution": institution,
                "project_type": project_type,
                "date": datetime.now().strftime("%B %d, %Y")
            }
        
        # Project abstract
        st.markdown("### Project Abstract")
        project_abstract = st.text_area(
            "Enter your project abstract",
            height=200,
            help="Provide a comprehensive summary of your project"
        )
        
        # Project specifics to improve content generation
        with st.expander("Additional Project Details (Improves Generation)"):
            st.text_area("Key Technologies Used", placeholder="e.g., Arduino Uno, HC-05 Bluetooth, L298N Motor Driver", key="technologies")
            st.text_area("Project Objectives", placeholder="List main objectives of your project", key="objectives")
            st.text_area("Project Components", placeholder="List main hardware/software components", key="components")
        
        # Generate documentation
        if st.button("Generate Documentation", type="primary"):
            if project_abstract:
                with st.spinner("Generating comprehensive documentation..."):
                    try:
                        st.session_state.doc_sections = {}
                        
                        # Generate each section
                        for section_key, section_info in SECTIONS.items():
                            # Skip cover page as it's generated directly
                            if section_key == 'COVER_PAGE':
                                continue
                                
                            with st.spinner(f"Generating {section_info['title']}..."):
                                # Use alternating models for different sections
                                model = groq_model if len(st.session_state.doc_sections) % 2 == 0 else google_model
                                
                                # Get additional context if available
                                technologies = st.session_state.get("technologies", "")
                                objectives = st.session_state.get("objectives", "")
                                components = st.session_state.get("components", "")
                                project_type = st.session_state.metadata.get("project_type", "")
                                
                                # Section-specific guidance
                                section_guidance = ""
                                if section_key == "CHAPTER_3":
                                    section_guidance = "Focus on hardware components like Arduino, sensors, and other electronic components. Describe their specifications and roles in detail."
                                elif section_key == "CHAPTER_4":
                                    section_guidance = "Explain the programming approach, include pseudocode or algorithm descriptions (not actual code), and detail the software workflow."
                                elif section_key == "LITERATURE_SURVEY":
                                    section_guidance = "Review at least 5-7 similar projects or related research papers, mentioning their approaches and results."
                                
                                prompt = f"""Generate the {section_info['title']} section for the following {project_type} project:
                                
                                Abstract: {project_abstract}
                                
                                Section Purpose: {section_info.get('description', '')}
                                
                                Additional Context:
                                - Technologies: {technologies}
                                - Objectives: {objectives}
                                - Components: {components}
                                - {section_guidance}
                                
                                Requirements:
                                1. This section should be approximately {section_info['pages']} pages long (about {section_info['pages'] * 500} words)
                                2. Use clear, straightforward explanations with appropriate technical depth
                                3. Organize content with clear sections and subsections
                                4. For headings, use plain text format (no special characters)
                                5. Write in a professional academic style with proper citations where appropriate
                                6. Use proper paragraph breaks for readability
                                7. Focus on factual information that would be relevant to this type of project
                                8. For CHAPTER_2, include descriptions of block diagrams (not the diagrams themselves)
                                9. Make this content unique, avoiding generic templates
                                10. Ensure all content is technically accurate and plausible
                                """
                                
                                if section_key == "CONTENTS":
                                    # Skip generating TOC as it's handled by Word
                                    st.session_state.doc_sections[section_key] = "TABLE OF CONTENTS"
                                else:
                                    response = model.invoke(prompt)
                                    st.session_state.doc_sections[section_key] = response.content
                        
                        # Create DOCX
                        create_docx(st.session_state.metadata)
                        st.session_state.has_generated = True
                        
                        # Add progress tracking
                        progress = st.progress(0)
                        for i in range(100):
                            progress.progress(i + 1)
                        
                        st.success("Documentation generated successfully!")
                        st.rerun()
                        
                    except Exception as e:
                        st.error(f"Error generating documentation: {str(e)}")
            else:
                st.warning("Please enter a project abstract.")

    with col2:
        if st.session_state.has_generated:
            st.subheader("Document Review & Feedback")
            
            # Section selection
            selected_section = st.selectbox(
                "Select section to review",
                options=[k for k in SECTIONS.keys() if k != "COVER_PAGE" and k != "CONTENTS"],
                format_func=lambda x: SECTIONS[x]['title']
            )
            
            if selected_section:
                st.markdown(f"### {SECTIONS[selected_section]['title']}")
                st.markdown(f"*{SECTIONS[selected_section].get('description', '')}*")
                st.markdown("---")
                st.markdown(st.session_state.doc_sections.get(selected_section, ""))
                
                # Feedback input
                feedback = st.text_area(
                    "Provide feedback for this section",
                    height=100,
                    placeholder="Explain desired changes or improvements..."
                )
                
                if st.button("Submit Feedback"):
                    with st.spinner("Updating section..."):
                        try:
                            # Alternate between models for feedback processing
                            model = google_model if len(st.session_state.feedback_history) % 2 == 0 else groq_model
                            
                            prompt = f"""Update the following documentation section based on user feedback:
                            
                            Current content:
                            {st.session_state.doc_sections[selected_section]}
                            
                            User feedback:
                            {feedback}
                            
                            Section Purpose: {SECTIONS[selected_section].get('description', '')}
                            
                            Requirements:
                            1. Maintain professional academic writing style
                            2. Implement the requested changes thoroughly
                            3. Keep the content within approximately {SECTIONS[selected_section]['pages'] * 500} words
                            4. Ensure the text remains technically accurate
                            5. Use proper paragraph structure and organization
                            6. Ensure content flows logically
                            7. Use plain text format for headings (no special characters)
                            """
                            
                            response = model.invoke(prompt)
                            
                            # Update section
                            st.session_state.doc_sections[selected_section] = response.content
                            
                            # Create updated DOCX
                            create_docx(st.session_state.metadata)
                            
                            # Record feedback
                            st.session_state.feedback_history.append({
                                "section": SECTIONS[selected_section]['title'],
                                "feedback": feedback,
                                "timestamp": datetime.now()
                            })
                            
                            st.success("Section updated successfully!")
                            st.rerun()
                            
                        except Exception as e:
                            st.error(f"Error updating section: {str(e)}")
            
            # Feedback history
            if st.session_state.feedback_history:
                with st.expander("Feedback History"):
                    for entry in reversed(st.session_state.feedback_history):
                        st.markdown(f"""
                        **Section:** {entry['section']}  
                        **Feedback:** {entry['feedback']}  
                        **Time:** {entry['timestamp'].strftime('%Y-%m-%d %H:%M:%S')}
                        ---
                        """)
            
            # Add document features info
            with st.expander("Document Features & Instructions"):
                st.markdown("""
                **Your document includes:**
                - Professionally formatted headings with borders
                - Placeholder for Table of Contents (with instructions to update in Word)
                - Consistent formatting throughout
                - Academic styling with Times New Roman font
                - 1.5 line spacing for improved readability
                
                **Important Notes:**
                1. **Page Borders**: To add page borders, open your document in Word, go to Design > Page Borders
                2. **Page Numbers**: The document contains placeholder text for page numbers. To add actual page numbers in Word:
                   - Go to Insert > Page Number > choose a position
                3. **Table of Contents**: Follow the instructions in the TOC section to generate an actual table of contents
                """)
            
            # Add download options
            st.download_button(
                "📥 Download Documentation",
                data=open("project_documentation.docx", "rb"),
                file_name="documentation.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

if __name__ == "__main__":
    main()